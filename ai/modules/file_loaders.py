import os
import csv
import pandas as pd
import docx
import PyPDF2
import markdown2
#from langchain.document_loaders import 
from langchain_core.documents import Document

class FileLoader:
    """
    Flexible file loader that determines file type by extension
    and loads content accordingly
    """
    # def document_format(text, metadata="ram"):
    #     for i in text:
    #         documents = [
    #             Document(page_content=text, metadata={"title": metadata}),
    #         ]
    #         return documents
    
    
    @staticmethod
    def load_file(file_path):
        """
        Load file based on its extension
        
        :param file_path: Full path to the file
        :return: Loaded file content (type depends on file type)
        """
        # Ensure file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        
        # Determine loader based on extension
        loaders = {
            '.csv': FileLoader._load_csv,
            '.txt': FileLoader._load_text,
            '.md': FileLoader._load_markdown,
            '.pdf': FileLoader._load_pdf,
            '.docx': FileLoader._load_docx
        }
        
        # Find appropriate loader
        loader = loaders.get(file_extension)
        if loader:
            return loader(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    @staticmethod
    def _load_csv(file_path):
        """Load CSV file"""
        try:
            return pd.read_csv(file_path)
        except Exception as e:
            print(f"Error loading CSV {file_path}: {e}")
            return None
    
    @staticmethod
    def _load_text(file_path):
        """Load plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error loading text file {file_path}: {e}")
            return None
    
    @staticmethod
    def _load_markdown(file_path):
        """Load Markdown file and convert to HTML"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return markdown2.markdown(file.read())
        except Exception as e:
            print(f"Error loading Markdown file {file_path}: {e}")
            return None
    
    @staticmethod
    def _load_pdf(file_path):
        """Extract text from PDF"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() + '\n'
                    
            return text
   
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")
            return None
    
    @staticmethod
    def _load_docx(file_path):
        """Extract text from Word document"""
        try:
            doc = docx.Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            print(f"Error loading Word document {file_path}: {e}")
            return None

# # Example usage
# def main():
#     # Example file paths (replace with your actual file paths)
#     files = [
#         '/path/to/your/document.csv',
#         '/path/to/your/notes.txt',
#         '/path/to/your/readme.md',
#         '/path/to/your/report.pdf',
#         '/path/to/your/document.docx'
#     ]
    
#     # Load and process each file
#     for file_path in files:
#         try:
#             # Attempt to load the file
#             content = FileLoader.load_file(file_path)
            
#             # Print file type and first few contents
#             print(f"\nFile: {file_path}")
#             if isinstance(content, pd.DataFrame):
#                 print("Type: CSV DataFrame")
#                 print(content.head())
#             elif isinstance(content, str):
#                 print("Type: Text-based File")
#                 print(content[:200])  # Print first 200 characters
        
#         except Exception as e:
#             print(f"Error processing {file_path}: {e}")

# if __name__ == "__main__":
#     main()
    
    
    
    
##README.md

# Load a specific file
# csv_content = FileLoader.load_file('/path/to/data.csv')
# pdf_text = FileLoader.load_file('/path/to/report.pdf')
if __name__ == '__main__':
    loader=FileLoader()
    file_path="../uploads/ram.pdf"
    data=loader.load_file(file_path)
    print(data)
    